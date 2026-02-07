from io import BytesIO
from docx import Document
from src.processing.extractors.base import BaseExtractor, ExtractedDocument


class DOCXExtractor(BaseExtractor):
    def extract(self, file_content: bytes) -> ExtractedDocument:
        """Extract text from a DOCX file using python-docx."""
        doc = Document(BytesIO(file_content))

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        content = "\n\n".join(paragraphs)

        return ExtractedDocument(
            content=content,
            page_count=None,  # DOCX doesn't have fixed pages
            metadata={"extractor": "python-docx"},
            method="native_docx",
        )
